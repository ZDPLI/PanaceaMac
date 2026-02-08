from __future__ import annotations

import re
from pathlib import Path


def extract_text(path: str) -> str:
    """Public API used by UI/RAG. Extract plain text from a supported file."""
    return extract_text_from_file(path)


def extract_text_from_file(path: str) -> str:
    """Backward-compatible alias."""
    p = Path(path)
    suffix = p.suffix.lower()

    if suffix in {".txt", ".md", ".log", ".csv", ".json"}:
        return _read_text(p)

    if suffix == ".docx":
        return _read_docx(p)

    if suffix == ".pdf":
        return _read_pdf(p)

    if suffix == ".rtf":
        return _read_rtf(p)

    # Fallback: try reading as text
    return _read_text(p)


def _read_text(p: Path) -> str:
    data = p.read_bytes()
    # Try UTF-8 first; then common Windows Cyrillic; finally latin-1 as a last resort.
    for enc in ("utf-8", "utf-8-sig", "cp1251", "cp866", "latin-1"):
        try:
            return data.decode(enc)
        except Exception:
            continue
    return data.decode("utf-8", errors="replace")


def _read_docx(p: Path) -> str:
    try:
        from docx import Document
    except Exception as e:
        raise RuntimeError("python-docx is required to read .docx") from e

    doc = Document(str(p))
    parts: list[str] = []

    # Paragraphs
    for para in doc.paragraphs:
        t = (para.text or "").strip()
        if t:
            parts.append(t)

    # Tables (often used in medical docs)
    for table in doc.tables:
        for row in table.rows:
            cells = []
            for cell in row.cells:
                t = (cell.text or "").strip()
                if t:
                    cells.append(t)
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts).strip()


def _read_pdf(p: Path) -> str:
    # Prefer pypdf if available; fall back to PyPDF2
    reader_cls = None
    exc = None
    try:
        import pypdf  # type: ignore
        reader_cls = pypdf.PdfReader
    except Exception as e:
        exc = e
    if reader_cls is None:
        try:
            import PyPDF2  # type: ignore
            reader_cls = PyPDF2.PdfReader
        except Exception as e:
            raise RuntimeError("pypdf or PyPDF2 is required to read .pdf") from (exc or e)

    parts: list[str] = []
    with p.open("rb") as f:
        reader = reader_cls(f)
        # Encrypted PDFs: attempt empty password (common) or fail gracefully
        try:
            is_encrypted = getattr(reader, "is_encrypted", False)
            if is_encrypted:
                try:
                    reader.decrypt("")  # type: ignore
                except Exception:
                    raise RuntimeError("PDF is encrypted and cannot be read without a password.")
        except Exception:
            pass

        pages = getattr(reader, "pages", [])
        for page in pages:
            try:
                txt = page.extract_text() or ""
            except Exception:
                txt = ""
            txt = txt.strip()
            if txt:
                parts.append(txt)

    return "\n\n".join(parts).strip()


def _read_rtf(p: Path) -> str:
    # Minimal RTF to text: strip control words and braces.
    raw = _read_text(p)
    text = re.sub(r"\\'[0-9a-fA-F]{2}", " ", raw)
    text = re.sub(r"\\[a-zA-Z]+-?\d* ?"," ", text)
    text = text.replace("{", " ").replace("}", " ")
    text = re.sub(r"\s+", " ", text)
    return text.strip()
